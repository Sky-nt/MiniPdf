"""
Generate 30 classic DOCX test files for the MiniPdf benchmark suite.

Each file tests a different Word document feature, from simple paragraphs
to tables, images, lists, headings, and mixed content.

Prerequisites:
    pip install python-docx Pillow

Usage:
    python generate_classic_docx.py            # output to ./output_docx/
    python generate_classic_docx.py --outdir /path/to/dir
"""

import argparse
import io
import os
import struct
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

OUTPUT_DIR = Path(__file__).parent / "output_docx"


def _create_test_png(width=120, height=80, color=(70, 130, 180)):
    """Create a minimal PNG in memory for embedding."""
    if HAS_PIL:
        img = PILImage.new("RGB", (width, height), color)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    # Fallback: build a minimal 1x1 PNG manually
    return _minimal_png(color)


def _minimal_png(color):
    """Build a minimal valid 1×1 PNG."""
    import zlib
    buf = io.BytesIO()
    buf.write(b"\x89PNG\r\n\x1a\n")
    # IHDR
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    _write_chunk(buf, b"IHDR", ihdr_data)
    # IDAT
    raw = b"\x00" + bytes(color)
    compressed = zlib.compress(raw)
    _write_chunk(buf, b"IDAT", compressed)
    # IEND
    _write_chunk(buf, b"IEND", b"")
    buf.seek(0)
    return buf


def _write_chunk(buf, chunk_type, data):
    import zlib
    buf.write(struct.pack(">I", len(data)))
    buf.write(chunk_type)
    buf.write(data)
    buf.write(struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF))


# ── Classic docx generators (1–30) ──────────────────────────────────────


def classic01_single_paragraph(path):
    """Single plain-text paragraph."""
    doc = Document()
    doc.add_paragraph("Hello, World! This is a simple single paragraph document created for benchmarking MiniPdf DOCX-to-PDF conversion.")
    doc.save(path)


def classic02_multiple_paragraphs(path):
    """Multiple plain-text paragraphs with implicit spacing."""
    doc = Document()
    for i in range(1, 6):
        doc.add_paragraph(f"This is paragraph {i}. It contains some sample text to test how MiniPdf handles multiple consecutive paragraphs with default spacing.")
    doc.save(path)


def classic03_headings(path):
    """Heading levels 1 through 4."""
    doc = Document()
    doc.add_heading("Heading Level 1", level=1)
    doc.add_paragraph("Content under heading 1.")
    doc.add_heading("Heading Level 2", level=2)
    doc.add_paragraph("Content under heading 2.")
    doc.add_heading("Heading Level 3", level=3)
    doc.add_paragraph("Content under heading 3.")
    doc.add_heading("Heading Level 4", level=4)
    doc.add_paragraph("Content under heading 4.")
    doc.save(path)


def classic04_bold_italic(path):
    """Bold, italic, and bold-italic text."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text. ")
    run_b = p.add_run("Bold text. ")
    run_b.bold = True
    run_i = p.add_run("Italic text. ")
    run_i.italic = True
    run_bi = p.add_run("Bold and italic text.")
    run_bi.bold = True
    run_bi.italic = True
    doc.save(path)


def classic05_font_sizes(path):
    """Various font sizes in one document."""
    doc = Document()
    for size in [8, 10, 12, 14, 18, 24, 36]:
        p = doc.add_paragraph()
        run = p.add_run(f"This text is {size}pt.")
        run.font.size = Pt(size)
    doc.save(path)


def classic06_font_colors(path):
    """Colored text runs."""
    doc = Document()
    colors = [
        ("Red text", RGBColor(255, 0, 0)),
        ("Green text", RGBColor(0, 128, 0)),
        ("Blue text", RGBColor(0, 0, 255)),
        ("Orange text", RGBColor(255, 165, 0)),
        ("Purple text", RGBColor(128, 0, 128)),
    ]
    for text, color in colors:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = color
    doc.save(path)


def classic07_alignment(path):
    """Paragraph alignment: left, center, right, justify."""
    doc = Document()
    lorem = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."

    p = doc.add_paragraph(lorem)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_run = p.runs[0] if p.runs else p.add_run("")
    doc.add_paragraph()

    p2 = doc.add_paragraph(lorem)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph(lorem)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p4 = doc.add_paragraph(lorem)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(path)


def classic08_bullet_list(path):
    """Bullet (unordered) list."""
    doc = Document()
    doc.add_heading("Shopping List", level=2)
    items = ["Apples", "Bananas", "Cherries", "Dates", "Elderberries"]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def classic09_numbered_list(path):
    """Numbered (ordered) list."""
    doc = Document()
    doc.add_heading("Steps to Success", level=2)
    steps = ["Define the goal", "Research the topic", "Create a plan", "Execute the plan", "Review results"]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.save(path)


def classic10_simple_table(path):
    """Simple 3x4 table with header row."""
    doc = Document()
    doc.add_heading("Employee Directory", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Department", "Email"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Alice Johnson", "Engineering", "alice@example.com"],
        ["Bob Smith", "Marketing", "bob@example.com"],
        ["Carol Williams", "Finance", "carol@example.com"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = val
    doc.save(path)


def classic11_table_with_shading(path):
    """Table with alternating row shading."""
    doc = Document()
    doc.add_heading("Quarterly Sales", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Quarter", "Revenue", "Expenses", "Profit"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    data = [
        ["Q1 2025", "$120,000", "$80,000", "$40,000"],
        ["Q2 2025", "$135,000", "$85,000", "$50,000"],
        ["Q3 2025", "$150,000", "$90,000", "$60,000"],
        ["Q4 2025", "$160,000", "$95,000", "$65,000"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            if ri % 2 == 0:
                _set_cell_shading(cell, "D9E2F3")
    doc.save(path)


def classic12_merged_cells_table(path):
    """Table with merged cells."""
    doc = Document()
    doc.add_heading("Schedule", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Time"
    table.rows[0].cells[1].text = "Monday"
    table.rows[0].cells[2].text = "Tuesday"
    table.rows[1].cells[0].text = "9:00 AM"
    merged = table.rows[1].cells[1].merge(table.rows[1].cells[2])
    merged.text = "Team Meeting"
    table.rows[2].cells[0].text = "10:00 AM"
    table.rows[2].cells[1].text = "Code Review"
    table.rows[2].cells[2].text = "Design Review"
    table.rows[3].cells[0].text = "2:00 PM"
    table.rows[3].cells[1].text = "Sprint Planning"
    table.rows[3].cells[2].text = "Retrospective"
    doc.save(path)


def classic13_long_document(path):
    """A longer document that spans multiple pages."""
    doc = Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph("This document is designed to span multiple pages to test pagination in MiniPdf.")
    for i in range(1, 16):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(
            f"This is section {i} of the report. It contains detailed analysis of the topic at hand. "
            "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. "
            "How vexingly quick daft zebras jump. The five boxing wizards jump quickly. "
            "Sphinx of black quartz, judge my vow." * 2
        )
    doc.save(path)


def classic14_mixed_content(path):
    """A document with headings, paragraphs, and a table mixed together."""
    doc = Document()
    doc.add_heading("Monthly Report", level=1)
    doc.add_paragraph("This report summarizes the key metrics for the month of January 2026.")

    doc.add_heading("Revenue Summary", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Amount"
    for i, (cat, amt) in enumerate([("Product Sales", "$85,000"), ("Services", "$42,000"), ("Subscriptions", "$28,000")]):
        table.rows[i + 1].cells[0].text = cat
        table.rows[i + 1].cells[1].text = amt

    doc.add_heading("Key Observations", level=2)
    doc.add_paragraph("Product sales increased by 15% compared to the previous quarter.")
    doc.add_paragraph("Service revenue remained stable with a slight upward trend.")

    doc.add_heading("Action Items", level=2)
    for item in ["Expand marketing campaign", "Hire two additional engineers", "Launch new subscription tier"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def classic15_indentation(path):
    """Paragraphs with various indentation levels."""
    doc = Document()
    doc.add_heading("Indentation Test", level=2)
    levels = [0, 36, 72, 108, 144]
    for pts in levels:
        p = doc.add_paragraph(f"This paragraph is indented by {pts} points from the left margin.")
        p.paragraph_format.left_indent = Pt(pts)
    # First-line indent
    p2 = doc.add_paragraph("This paragraph has a first-line indent of 36 points. The remaining lines wrap normally back to the left margin.")
    p2.paragraph_format.first_line_indent = Pt(36)
    doc.save(path)


def classic16_line_spacing(path):
    """Different line spacing options."""
    doc = Document()
    doc.add_heading("Line Spacing Test", level=2)
    text = "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. How vexingly quick daft zebras jump."
    from docx.shared import Pt as PtShared
    for spacing_val, label in [(1.0, "Single"), (1.5, "1.5 Lines"), (2.0, "Double")]:
        doc.add_paragraph(f"{label} spacing:")
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = spacing_val
        doc.add_paragraph()
    doc.save(path)


def classic17_page_break(path):
    """Document with explicit page breaks."""
    doc = Document()
    doc.add_heading("Page 1", level=1)
    doc.add_paragraph("Content on the first page.")
    doc.add_page_break()
    doc.add_heading("Page 2", level=1)
    doc.add_paragraph("Content on the second page after a page break.")
    doc.add_page_break()
    doc.add_heading("Page 3", level=1)
    doc.add_paragraph("Content on the third page.")
    doc.save(path)


def classic18_embedded_image(path):
    """Document with an embedded PNG image."""
    doc = Document()
    doc.add_heading("Image Test", level=2)
    doc.add_paragraph("Below is an embedded test image:")
    img_buf = _create_test_png(200, 100, (70, 130, 180))
    doc.add_picture(img_buf, width=Inches(3))
    doc.add_paragraph("The image above should be a blue rectangle.")
    doc.save(path)


def classic19_multiple_images(path):
    """Multiple images at different sizes."""
    doc = Document()
    doc.add_heading("Multiple Images", level=2)
    colors = [(220, 50, 50), (50, 180, 50), (50, 50, 220)]
    names = ["Red", "Green", "Blue"]
    for color, name in zip(colors, names):
        doc.add_paragraph(f"{name} image:")
        img_buf = _create_test_png(160, 80, color)
        doc.add_picture(img_buf, width=Inches(2.5))
    doc.save(path)


def classic20_table_with_many_rows(path):
    """Large table with 20 data rows."""
    doc = Document()
    doc.add_heading("Product Catalog", level=2)
    table = doc.add_table(rows=21, cols=4)
    table.style = "Table Grid"
    headers = ["ID", "Product", "Category", "Price"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    products = [
        ("Laptop", "Electronics", "$999"),
        ("Mouse", "Accessories", "$29"),
        ("Keyboard", "Accessories", "$59"),
        ("Monitor", "Electronics", "$349"),
        ("Headphones", "Audio", "$149"),
        ("Webcam", "Electronics", "$79"),
        ("USB Hub", "Accessories", "$25"),
        ("Desk Lamp", "Office", "$45"),
        ("Chair", "Furniture", "$299"),
        ("Standing Desk", "Furniture", "$599"),
        ("Printer", "Electronics", "$199"),
        ("Scanner", "Electronics", "$129"),
        ("Router", "Networking", "$89"),
        ("Cable Kit", "Accessories", "$19"),
        ("Mousepad", "Accessories", "$15"),
        ("Surge Protector", "Electronics", "$35"),
        ("External SSD", "Storage", "$109"),
        ("Flash Drive", "Storage", "$12"),
        ("Drawing Tablet", "Electronics", "$249"),
        ("Microphone", "Audio", "$179"),
    ]
    for ri, (prod, cat, price) in enumerate(products):
        table.rows[ri + 1].cells[0].text = str(ri + 1)
        table.rows[ri + 1].cells[1].text = prod
        table.rows[ri + 1].cells[2].text = cat
        table.rows[ri + 1].cells[3].text = price
    doc.save(path)


def classic21_nested_lists(path):
    """Nested bullet list (simulated with indentation)."""
    doc = Document()
    doc.add_heading("Project Structure", level=2)
    items = [
        (0, "src/"),
        (1, "MiniPdf/"),
        (2, "MiniPdf.cs"),
        (2, "PdfDocument.cs"),
        (2, "PdfWriter.cs"),
        (1, "MiniPdf.Tests/"),
        (2, "DocxToPdfConverterTests.cs"),
        (0, "scripts/"),
        (1, "Run-Benchmark.ps1"),
        (0, "README.md"),
    ]
    for level, text in items:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Pt(18 * level)
    doc.save(path)


def classic22_horizontal_rule(path):
    """Document with horizontal line separators (simulated)."""
    doc = Document()
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Content for section A goes here with enough text to see the layout.")
    _add_horizontal_rule(doc)
    doc.add_heading("Section B", level=2)
    doc.add_paragraph("Content for section B goes here below the horizontal divider.")
    _add_horizontal_rule(doc)
    doc.add_heading("Section C", level=2)
    doc.add_paragraph("Final section content.")
    doc.save(path)


def classic23_mixed_formatting_runs(path):
    """Single paragraph with many formatting runs."""
    doc = Document()
    doc.add_heading("Mixed Formatting", level=2)
    p = doc.add_paragraph()
    p.add_run("Normal, ")
    r1 = p.add_run("BOLD, ")
    r1.bold = True
    r2 = p.add_run("italic, ")
    r2.italic = True
    r3 = p.add_run("large, ")
    r3.font.size = Pt(18)
    r4 = p.add_run("small, ")
    r4.font.size = Pt(8)
    r5 = p.add_run("RED, ")
    r5.font.color.rgb = RGBColor(255, 0, 0)
    r6 = p.add_run("underlined.")
    r6.underline = True
    doc.save(path)


def classic24_two_column_table_layout(path):
    """Two-column layout using a borderless table."""
    doc = Document()
    doc.add_heading("Two-Column Layout", level=2)
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.text = (
        "Left column content. This is the first column of a two-column layout. "
        "It demonstrates how tables can be used for text layout purposes."
    )
    right_cell.text = (
        "Right column content. This is the second column. "
        "Both columns should render side-by-side in the PDF output."
    )
    doc.save(path)


def classic25_title_and_subtitle(path):
    """Title page with title and subtitle styles."""
    doc = Document()
    doc.add_paragraph("MiniPdf Benchmark Report", style="Title")
    doc.add_paragraph("Automated DOCX-to-PDF Conversion Quality Assessment", style="Subtitle")
    doc.add_paragraph()
    doc.add_paragraph("Prepared by: MiniPdf Team")
    doc.add_paragraph("Date: March 2026")
    doc.add_page_break()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This document tests the Title and Subtitle styles in MiniPdf conversion.")
    doc.save(path)


def classic26_table_alignment(path):
    """Table with aligned cell content (left, center, right)."""
    doc = Document()
    doc.add_heading("Cell Alignment Test", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = [("Left", WD_ALIGN_PARAGRAPH.LEFT),
               ("Center", WD_ALIGN_PARAGRAPH.CENTER),
               ("Right", WD_ALIGN_PARAGRAPH.RIGHT)]
    for ci, (h, align) in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].alignment = align
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ["Alice", "Engineering", "$95,000"],
        ["Bob", "Marketing", "$82,000"],
        ["Carol", "Finance", "$88,000"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            cell.paragraphs[0].alignment = headers[ci][1]
    doc.save(path)


def classic27_long_paragraph_wrapping(path):
    """Very long paragraph to test word wrapping."""
    doc = Document()
    doc.add_heading("Word Wrapping Test", level=2)
    long_text = (
        "This is a very long paragraph designed to test how MiniPdf handles word wrapping across "
        "line boundaries. The text should flow naturally from one line to the next without any "
        "awkward breaks or overflow. " * 10
    )
    doc.add_paragraph(long_text)
    doc.save(path)


def classic28_special_characters(path):
    """Document with special characters and symbols."""
    doc = Document()
    doc.add_heading("Special Characters", level=2)
    doc.add_paragraph("Ampersand: &, Less-than: <, Greater-than: >, Quotes: \" '")
    doc.add_paragraph("Copyright: \u00a9, Registered: \u00ae, Trademark: \u2122")
    doc.add_paragraph("Em-dash: \u2014, En-dash: \u2013, Ellipsis: \u2026")
    doc.add_paragraph("Currency: $ \u20ac \u00a3 \u00a5")
    doc.add_paragraph("Math: \u00b1 \u00d7 \u00f7 \u2260 \u2264 \u2265 \u221e")
    doc.save(path)


def classic29_table_with_image(path):
    """Table with an image in one cell."""
    doc = Document()
    doc.add_heading("Product Card", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Product"
    table.rows[0].cells[1].text = "Description"
    # Add image to first column of second row
    img_buf = _create_test_png(100, 60, (34, 139, 34))
    table.rows[1].cells[0].paragraphs[0].add_run().add_picture(img_buf, width=Inches(1.5))
    table.rows[1].cells[1].text = "MiniPdf Widget - A compact, efficient tool for PDF conversion. Lightweight and dependency-free."
    doc.save(path)


def classic30_comprehensive_report(path):
    """A comprehensive document combining many features."""
    doc = Document()
    # Title
    doc.add_paragraph("Annual Technology Report 2026", style="Title")
    doc.add_paragraph("A Comprehensive Overview", style="Subtitle")
    doc.add_page_break()

    # Table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    for i, title in enumerate(["Executive Summary", "Market Analysis", "Technology Trends", "Financial Overview", "Recommendations"], 1):
        doc.add_paragraph(f"{i}. {title}")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides a comprehensive analysis of the technology landscape in 2026. "
        "Key findings include continued growth in AI adoption, increased focus on sustainability, "
        "and emerging trends in quantum computing."
    )

    # Market Analysis with table
    doc.add_heading("2. Market Analysis", level=1)
    doc.add_paragraph("The following table summarizes market share across key sectors:")
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Sector", "Market Share", "Growth"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    sectors = [
        ("Cloud Computing", "34%", "+12%"),
        ("AI/ML", "28%", "+23%"),
        ("Cybersecurity", "22%", "+18%"),
        ("IoT", "16%", "+8%"),
    ]
    for ri, (sector, share, growth) in enumerate(sectors):
        table.rows[ri + 1].cells[0].text = sector
        table.rows[ri + 1].cells[1].text = share
        table.rows[ri + 1].cells[2].text = growth

    # Technology Trends
    doc.add_heading("3. Technology Trends", level=1)
    doc.add_paragraph("Key trends identified:")
    for trend in ["Generative AI integration in enterprise software",
                  "Edge computing for real-time processing",
                  "Green technology and sustainable computing",
                  "Zero-trust security architectures",
                  "Low-code/no-code platform expansion"]:
        doc.add_paragraph(trend, style="List Bullet")

    # Image section
    doc.add_heading("4. Visual Summary", level=2)
    doc.add_paragraph("Growth indicator chart (placeholder):")
    img_buf = _create_test_png(300, 150, (46, 84, 150))
    doc.add_picture(img_buf, width=Inches(4))

    # Recommendations
    doc.add_heading("5. Recommendations", level=1)
    for i, rec in enumerate(["Invest in AI-driven automation tools",
                              "Prioritize cloud-native architectures",
                              "Strengthen cybersecurity posture",
                              "Explore quantum computing partnerships"], 1):
        doc.add_paragraph(rec, style="List Number")

    doc.save(path)


# ── Helpers ──────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color):
    """Set a table cell's background shading."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shd)


def _add_horizontal_rule(doc):
    """Add a paragraph that looks like a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "auto",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Main ─────────────────────────────────────────────────────────────────

ALL_GENERATORS = [
    classic01_single_paragraph,
    classic02_multiple_paragraphs,
    classic03_headings,
    classic04_bold_italic,
    classic05_font_sizes,
    classic06_font_colors,
    classic07_alignment,
    classic08_bullet_list,
    classic09_numbered_list,
    classic10_simple_table,
    classic11_table_with_shading,
    classic12_merged_cells_table,
    classic13_long_document,
    classic14_mixed_content,
    classic15_indentation,
    classic16_line_spacing,
    classic17_page_break,
    classic18_embedded_image,
    classic19_multiple_images,
    classic20_table_with_many_rows,
    classic21_nested_lists,
    classic22_horizontal_rule,
    classic23_mixed_formatting_runs,
    classic24_two_column_table_layout,
    classic25_title_and_subtitle,
    classic26_table_alignment,
    classic27_long_paragraph_wrapping,
    classic28_special_characters,
    classic29_table_with_image,
    classic30_comprehensive_report,
]


def main():
    parser = argparse.ArgumentParser(description="Generate classic DOCX test files")
    parser.add_argument("--outdir", default=str(OUTPUT_DIR), help="Output directory")
    args = parser.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    print(f"Generating {len(ALL_GENERATORS)} classic DOCX files to {outdir}/")
    print()

    passed = 0
    failed = 0
    for i, gen_func in enumerate(ALL_GENERATORS, 1):
        name = gen_func.__name__
        filename = f"docx_{name}.docx"
        filepath = outdir / filename
        try:
            gen_func(str(filepath))
            size_kb = filepath.stat().st_size / 1024
            print(f"  OK  {filename} ({size_kb:.1f} KB)")
            passed += 1
        except Exception as e:
            print(f"  ERR {filename}: {e}")
            failed += 1

    print(f"\nDone! Passed: {passed}, Failed: {failed}, Total: {len(ALL_GENERATORS)}")


if __name__ == "__main__":
    main()
